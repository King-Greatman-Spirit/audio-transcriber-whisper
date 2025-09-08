import sys
import os
import whisper
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def transcribe_audio(audio_path, output_folder="docx"):
    # Load model (base for speed, large for accuracy)
    model = whisper.load_model("base")

    print(f"🔄 Transcribing {audio_path} ... this may take a while.")
    result = model.transcribe(audio_path)

    # Extract text
    transcript = result["text"]

    # Ensure output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    base_name = os.path.splitext(os.path.basename(audio_path))[0]

    # Save as .txt (inside docx/)
    txt_file = os.path.join(output_folder, base_name + ".txt")
    with open(txt_file, "w", encoding="utf-8") as f:
        f.write(transcript)
    print(f"✅ Transcript saved as TXT: {txt_file}")

    # Save as .docx (inside docx/)
    docx_file = os.path.join(output_folder, base_name + ".docx")
    doc = Document()
    doc.add_heading("Audio Transcription", level=1)
    doc.add_paragraph(transcript)
    doc.save(docx_file)
    print(f"✅ Transcript saved as Word DOCX: {docx_file}")

    # Save as .pdf (inside docx/)
    pdf_file = os.path.join(output_folder, base_name + ".pdf")
    c = canvas.Canvas(pdf_file, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica", 11)

    # Write transcript with word wrapping
    y = height - 50
    for line in transcript.split("\n"):
        for chunk in [line[i:i+90] for i in range(0, len(line), 90)]:
            c.drawString(50, y, chunk)
            y -= 15
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - 50
    c.save()

    print(f"✅ Transcript saved as PDF: {pdf_file}")

    return transcript


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("❌ Usage: python audioToText.py <audiofile>")
    else:
        audio_file = sys.argv[1]
        transcribe_audio(audio_file)
